
from dotenv import load_dotenv
import os
import pandas as pd
import matplotlib.pyplot as plt
import docx
import requests

load_dotenv()  # This method will load the .env file

openai_api_key = os.getenv('OPENAI_API_KEY')  # Now you can access the API key
from openai import OpenAI
import tiktoken
client = OpenAI(api_key=openai_api_key)
embedding_model = "text-embedding-3-large"
gpt4 = "gpt-4-0125-preview"
gpt3 = "gpt-3.5-turbo-0125"

from docx import Document

# Load the Word documents
performance_task_description_doc = Document('files/Data Analyst Performance Task_v2.docx')
performance_task_example_doc = Document('files/Sample Memo - BZA Fees-2.docx')




# Read the text of the document
performance_task_description_text = []
for paragraph in performance_task_description_doc.paragraphs:
    performance_task_description_text.append(paragraph.text)

performance_task_example_text = []
for paragraph in performance_task_example_doc.paragraphs:
    performance_task_example_text.append(paragraph.text)

# Join the text into a single string
performance_task_description = '\n'.join(performance_task_description_text)
performance_task_example = '\n'.join(performance_task_example_text)

def classify_column(row):
    code = row['Violation Cited (State/Local Code Reference)']
    if code.startswith('154:383'):
        return 257.50, 'Careless operation of a vehicle', 'local'
    elif code.startswith('32:58'):
        return 302.50, 'Careless operation of a vehicle', 'state'
    elif code.startswith('154:303'):
        return 207.50, 'D/L not on person', 'local'
    elif code.startswith('32:411.1'):
        return 302.50, 'D/L not on person', 'state'
    elif code.startswith('154:236'):
        return 227.50, 'Disregard red light', 'local'
    elif code.startswith('32:232'):
        return 302.50, 'Disregard red light', 'state'
    elif code.startswith('154:400'):
        return 207.50, 'Drivers to use reasonable vigilance', 'local'
    elif code.startswith('154:401'):
        return 227.50, 'Following too closely', 'local'
    elif code.startswith('32:81'):
        return 302.50, 'Following too closely', 'state'
    elif code.startswith('154:240'):
        return 157.50, 'Improper lane use', 'local'
    elif code.startswith('32:79'):
        return 302.50, 'Improper lane use', 'state'
    elif code.startswith('154:436'):
        return 157.50, 'Improper turn', 'local'
    elif code.startswith('32:101'):
        return 302.50, 'Improper turn', 'state'
    elif code.startswith('154:307'):
        return 157.50, 'License plates', 'local'
    elif code.startswith('32:53A'):
        return 302.50, 'License plates', 'state'
    elif code.startswith('154:1298'):
        return 207.50, 'No/expired brake tag', 'local'
    elif code.startswith('32:53D'):
        return 302.50, 'No/expired brake tag', 'state'
    elif code.startswith('154:304'):
        return 607.50, 'Unlawful use of license', 'local'
    elif code.startswith('32:414.1'):
        return 627.50, 'Unlawful use of license', 'state'
    elif code.startswith('154:482'):
        return 227.50, 'Yield right of way', 'local'
    else:
        return 302.50, 'Yield right of way', 'state'


def chart_citations(ws, chart_title):
    # Gather data from 'ws' into a list of lists (rows)
    data_rows = []
    for row in ws.iter_rows(values_only=True):
        data_rows.append(list(row))

    # The first row typically contains the column headers
    column_headers = data_rows.pop(0)  # This removes the first row and saves it as column headers

    # Create a DataFrame from the rows of data
    df = pd.DataFrame(data_rows, columns=column_headers)

    # Apply the function to each row and create a new DataFrame with the results
    classifications = df.apply(classify_column, axis=1)

    # The result is a Series of tuples. We can convert this Series into a new DataFrame
    classifications_df = pd.DataFrame(classifications.tolist(), columns=['Revenue', 'Violation Type', 'Scope'])

    # Now, concatenate this new DataFrame with the original 'df'
    df = pd.concat([df, classifications_df], axis=1)

    # Convert 'Date' column to datetime
    df['Violation Date'] = pd.to_datetime(df['Violation Date'])

    # Calculate the number of months the data spans over
    """num_months = (df['Violation Date'].dt.to_period('M').max() - df['Violation Date'].dt.to_period('M').min()).n + 1
    print(f"The data spans over {num_months} months.")"""

    # Add a column for month-year
    df['Month-Year'] = df['Violation Date'].dt.to_period('M')

    # Group by 'Violation Type' and 'Scope', and count the number of citations
    total_violations_by_type = df.groupby(['Violation Type', 'Scope'])['Violation Type'].count().unstack()

    # Plotting the total number of violations per violation type
    plt.figure(figsize=(10, 5))  # Adjusts the size of the plot

    # Use kind='bar' to create a bar chart with separate bars for 'State' and 'Local'
    total_violations_by_type.plot(kind='bar', stacked=False)

    # Add titles and labels
    plt.title('Number of Citations by Violation Type')
    plt.xlabel('Violation Type')
    plt.ylabel('Number of Citations')

    # Shorten or rotate x-axis labels for better readability
    plt.xticks(rotation=45, ha='right')  # Rotate the labels and align them right for better fit

    # Add legend to distinguish between 'State' and 'Local' totals
    plt.legend(title='Scope')

    # Save the plot
    plt.savefig(f'files/{chart_title}_citations_plot.png', bbox_inches='tight')

    plt.savefig(f'files/{chart_title}_citations_plot.png', bbox_inches='tight')

    # Group by 'Month-Year' and 'Scope', then calculate the total revenue
    total_revenue = df.groupby(['Violation Type', 'Scope'])['Revenue'].sum().unstack()

    total_revenue.plot(kind='bar', stacked=False)  # Use kind='bar' to create a bar chart
    # total_monthly_revenue['local'].plot(kind='line', marker='o', label='Total Local Revenue')
    # total_monthly_revenue['state'].plot(kind='line', marker='x', label='Total State Revenue')

    # Add titles and labels
    plt.title(f'{chart_title} Citation Revenue by Violation Type')
    plt.xlabel('Violation Type')
    plt.ylabel('Total Revenue ($)')

    # Rotate date labels for better readability
    plt.xticks(rotation=45, ha='right')

    # Add legend to distinguish between local and state totals
    plt.legend()

    plt.savefig(f'files/{chart_title}_revenue_plot.png', bbox_inches='tight')

    return df



def get_embedding(text, model="text-embedding-3-large"):
    print(text)
    text = text.replace("\n", " ")
    emb = client.embeddings.create(input = [text], model=model).data[0].embedding
    df = pd.DataFrame({'text': text, 'embeddings': emb})
    return df



def combine_embeddings(dfs):
    combined_df = pd.concat(dfs, ignore_index=True)
    print(combined_df)
    return combined_df


def save_embedding(dfs, model):

    combined_embedding = combine_embeddings(dfs)

    #emb = client.embeddings.create(input = f"{combined_embedding}", model=model).data[0].embedding
    combined_embedding.to_csv('downloads/embedded_text.csv', index=False)


def load_embedding():
    df = pd.read_csv('downloads/embedded_text.csv')
    return df


def get_text_from_python(file_path):
    # Open the Python file in read mode
    with open(file_path, 'r') as file:
        # Read the contents of the file
        raw_text = file.read()
    return raw_text


def get_text_from_word(file_path):
    # Load the Word document
    doc = docx.Document(file_path)

    # Extract text from the document
    raw_text = ''
    for paragraph in doc.paragraphs:
        raw_text += paragraph.text + '\n'

    return raw_text



def get_text_from_ws(worksheet):
    # Extract text from each cell and concatenate it into a single string
    text = ''
    for row in worksheet.iter_rows(values_only=True):
        for cell in row:
            if cell is not None:
                text += str(cell) + ' '
    return text


def get_months(cell_value):
    completion = client.chat.completions.create(
        model=gpt4,
        messages=[
            {"role": "system", "content": "you will return only the number of months implied by given text."
                                          "no need to put months after the text, just the number. "},
            {"role": "user", "content": cell_value}
        ]
    )
    return completion.choices[0].message


def get_summarized_context(prompt, model, text):
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "you are trying to reduce token usage by sumarizing relvant parts of a given text. "
                                          "the user will provide you with a prompt that will be answered in another api request. "
                                          "your job is to provide the content of the file given below that is relevant to the question being asked. "
                                          "and return as much relevant information as possible so that the question can be answered in another request. "
                                          "Dont answer the question, just return the parts of the text that are relevant "
                                          "to the question, along with supporting information. Provide as many specific details examples, and quotations as you can. "
                                          "Try to shoot for about 500 tokens if possible but feel free to use more or less. \n"
                                          f"file: {text}"},
            {"role": "user", "content": prompt}
        ]
    )
    return completion.choices[0].message

def is_question_relevant(prompt, model, text):
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "you will check if the users question could be interpreted to be about "
                                          "the project, which is a "
                                          "telegram bot that generates memos about new orleans trafic citations "
                                          "and revenue implications, and answers questions about these memos, "
                                          "the python code that generates them, "
                                          "and the data and analyses used in the project."
                                          "if the users prompt is related to the project respond True, if not "
                                          "respond False. You may only respond with True or False.  "
                                          "Below is a summary of one of the memos for reference\n"
                                          f"file: {text}"},
            {"role": "user", "content": prompt}
        ]
    )
    return completion.choices[0].message


def get_completion(prompt, model, text, math_sheet, table, messager, methods, analysis):

    memo_summary = get_summarized_context(prompt, model, text)

    relevant = is_question_relevant(prompt, gpt3, memo_summary.content)
    print(f"Is the question relevant to the project? {relevant}")

    if relevant.content == "True":

        math_sheet_summary = get_summarized_context(prompt, gpt3, math_sheet)
        math_sheet_summary = math_sheet_summary.content

        table_summary = get_summarized_context(prompt, gpt3, table)
        table_summary = table_summary.content

        messager_summary = get_summarized_context(prompt, gpt3, messager)
        messager_summary = messager_summary.content

        methods_summary = get_summarized_context(prompt, gpt3, methods)
        methods_summary = methods_summary.content

        analysis_summary = get_summarized_context(prompt, gpt3, analysis)
        analysis_summary = analysis_summary.content

    else:
        text = memo_summary.content
        math_sheet_summary = ""
        table_summary = ""
        messager_summary = ""
        methods_summary = ""
        analysis_summary = ""

    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a data scientist, you respond in a casual to the point manner. "
                                          "You are working on a data analysis project for the city of New Orleans. "
                                          "You are analyzing the effects of New Orleans Police Department Traffic Citations "
                                          "on the annual budget. "
                                          "Answer questions about the memo, math sheet, and table provided in a "
                                          "slightly casual, but professional manner. "
                                          "You also answer questions about the methods analysis, and telegram bot used in the project. "
                                          "the code used for the methods and analysis is in the python files provided. "
                                          "feel free to provide examples but dont reveal any api keys."
                                          "You can also answer questions about the telegram bot used to communicate with you. "
                                          "if the answer to a question isn't related to the project, "
                                          "just say so, and then answer the question as best you can anyways. "
                                          "Below are summarized Project Details specific to the question: \n"
                                          f"Memo: {text} \n"
                                          f"Math Sheet: {math_sheet_summary} \n"
                                          f"Table: {table_summary} \n"
                                          f"Messager: {messager_summary} \n"
                                          f"Methods: {methods_summary} \n"
                                          f"Analysis: {analysis_summary} \n"
            },
            {"role": "user", "content": prompt}
        ]
    )
    # Retrieve the number of tokens used in the completion

    completion_tokens = completion.usage.completion_tokens
    completion_price = completion_tokens * (0.03/1000)
    print(f"Number of completion tokens: {completion_tokens} ${completion_price}\n")

    prompt_tokens = completion.usage.prompt_tokens
    prompt_price = prompt_tokens * (0.01/1000)
    print(f"Number of prompt tokens: {prompt_tokens} ${prompt_price}\n")

    tokens_used = completion.usage.total_tokens
    total_price = completion_price + prompt_price
    print(f"\nNumber of tokens used: {tokens_used}\nTotal Price: ${total_price}\n")
    return completion.choices[0].message


def get_memo_completion(prompt, violation_types, totals, results, model, previous_text):
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a data scientist text completion bot for the city of new orleans, "
                                          "skilled in analyzing data and conveying important findings concisely. "
                                          "You respond in first person in a casual, easily understandable, but professional manner. "
                                          "Dont use word that are too complex, and keep the memo simple and to the point. "
                                          "You are writing a memo to your manager about the findings of your analysis "
                                          "of effects of New Orleans Police Department Traffic Citations "
                                          "on the annual budget. The memo should be clear, concise, and informative."
                                          "The memo will demonstrate your ability to communicate, "
                                          "complex data in a simple and understandable way. "
                                          "Remember that the task is the analyze the NOPD Citation data, "
                                          "and that the columns in the dataset that are not "
                                          "specified as nopd were submitted by state police, "
                                          "and therefore should not be considered as NOPD citations in the analysis. "
                                          "The user will indicate which part of the prompt that needs to be completed. "
                                          "The formatting of the memo is already completed, "
                                          "so you must always leave out all section titles, additional formatting, "
                                          "and asterisks like **...**. "
                                          "Additionally since your purpose is to complete the memo section by section, "
                                          "the sections that are previously written have been provided below: \n"
                                          f"{previous_text}.\n "
                                          f"Do not include previous_text in the response, it is only for "
                                          f"reference so the final memo is not redundant. "
                                          "Make sure the memo in its entirety is coherent, flows well, and isn't redundant. "
                                          "No need to label the section with a title such as BACKGROUND, "
                                          "as that has already been accounted for"
                                          "The memo will be based off of the following prompt:"
                                          f"{performance_task_description}.\n"
                                          "The memo will be based off of the following data:"
                                          f"{violation_types} contains the number of citations and revenue for each violation type.\n"
                                          f"{totals} contains the total number of citations and the total revenue. \n"
                                          f"{results} contains the required budget impact analysis. "
                                          "Make sure to discuss the budget impact in detail\n"
                                          "The memo will be based off of the following example, but dont include "
                                          "the to: from: section or any formatting, or chat-like responses"
                                          "as the final product may end up being redundant:"
                                          f"{performance_task_example}.\n"},
            {"role": "user", "content": prompt}
        ]
    )
    print(completion.choices[0].message.content)
    return completion.choices[0].message

