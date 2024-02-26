import os
import logging
from dotenv import load_dotenv

# Import the pandas library
import pandas as pd

# Import openpyxl and the column_index_from_string function
from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string
from openpyxl.utils.dataframe import dataframe_to_rows

# Import python-docx
from docx import Document
from docx.shared import Pt, Inches

# Import the methods from the methods.py file
import methods
import analysis
import embeddings

import telegram
from telegram import Update, Bot
from telegram.ext import Application, CommandHandler, ContextTypes, MessageHandler, filters

# Create a logger
logger = logging.getLogger(__name__)

# Set the level for the logger
logger.setLevel(logging.INFO)

# Create a file handler
handler = logging.FileHandler('logfile.log')

# Set the formatter for the handler
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)

# Add the handler to the logger
logger.addHandler(handler)

load_dotenv()  # This method will load the .env file

telegram_api_key = os.getenv('TELEGRAM_API_KEY')
bot = Bot(token=telegram_api_key)
chat_id = '6310217725'


print('Messaging Filters Configured')
print('------\nListening For Messages...\n------\n')
memo_text = analysis.get_text_from_word('files/memo.docx')
math_sheet = analysis.get_text_from_word('files/Math_Calculations_Work.docx')

# Load the Excel file
wb = load_workbook('files/excel.xlsx')
summary = wb['Summary']
table = analysis.get_text_from_ws(summary)

mes_text = analysis.get_text_from_python('messager.py')
methods_text = analysis.get_text_from_python('methods.py')
analysis_text = analysis.get_text_from_python('analysis.py')

logger.info('Files Loaded')
logger.info('\n------\nListening For Messages...\n------\n')

# Define a `/start` command handler.
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    launch_user = update.message.from_user
    print(launch_user)
    message_text = f'{launch_user.first_name} {launch_user.last_name} has launched the bot. (@{launch_user.username})'

    await bot.send_message(chat_id=chat_id, text=message_text)
    text = (
        f"Hello, I am a bot that generates memos for the City of New Orleans Office of Performance and Accountability. "
        f"Right now I am configured to produce Traffic Citation Revenue Analysis Memos detailed in the document below. "
        f"Use the command /get_memo to generate a memo based on the sample data.  "
        f"Or send me an excel file with the same format as the sample and "
        f"I will fill the missing cells generate a memo for you "
        f"based on the data it contains (it might take me a couple of minutes fyi).  "
        f"Also, feel free to ask me any questions. "
        f"\n\nPlease note that I am still in development and may not always work as expected. "
    )
    await update.message.reply_html(text=text)

    word_document_path = 'files/Data Analyst Performance Task_v2.docx'
    # Use the send_document method to send the document
    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(word_document_path, 'rb'))

    excel_document_path = '/Users/alexanderbauer/Desktop/Cataphora/cityTask/files/' \
                          'CLEAN_Innovation Manager Performance Task - Sample Data.xlsx'
    # Use the send_document method to send the document
    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(excel_document_path, 'rb'))
    logger.info(f'Bot Started by {launch_user.first_name} {launch_user.last_name}')


# Define a `/start` command handler.
async def get_memo(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    """Send a message with a button that opens the web app."""
    launch_user = update.message.from_user
    print(f"From: {launch_user}")
    text = (
        f"I'm working on generating a memo based on the sample data, covering traffic citation revenue "
        f"for the City of New Orleans from August 2017 to August 2018. "
        f"I'll send it over soon."
    )
    await update.message.reply_html(text=text)
    download_path = '/Users/alexanderbauer/Desktop/Cataphora/cityTask/files/' \
                    'CLEAN_Innovation Manager Performance Task - Sample Data.xlsx'
    generate_memo(download_path,
                  f'files/memo.docx',
                  f'files/excel.xlsx')
    await update.message.reply_text('I have successfully generated the memo:')
    print(f"I have successfully generated the memo:")
    await context.bot.send_document(chat_id=update.effective_chat.id,
                                    document=open(f'files/memo.docx', 'rb'))
    await context.bot.send_document(chat_id=update.effective_chat.id,
                                    document=open(f'files/excel.xlsx', 'rb'))
    await context.bot.send_document(chat_id=update.effective_chat.id,
                                    document=open(f'files/Math_Calculations_Work.docx', 'rb'))

    logger.info(f'Memo Generated by {launch_user.first_name} {launch_user.last_name}')


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    from_user = update.message.from_user
    print(f"From: {from_user}")
    # print(f"From ID: {from_user.id}")

    # Check if there's text in the message
    received_text = update.message.text
    if received_text:
        print(f"Received Message: {received_text}")
        if received_text.startswith("/"):
            return
        else:
            try:
                #results = embeddings.ask(received_text, print_message=True)
                results = analysis.get_completion(received_text, analysis.gpt4, memo_text, math_sheet,
                                                  table, mes_text, methods_text, analysis_text)
                logger.info(f"GPT4 Message: {results}")
            except Exception as e:
                print(f"Error: GPT4 Failed... {e}")
                logger.error(f"Error: GPT4 Failed... {e}")
                results = analysis.get_completion(received_text, analysis.gpt3, memo_text, math_sheet,
                                                  table, mes_text, methods_text, analysis_text)
                logger.info(f"GPT3 Message: {results}")

            print(f"Results: {results}")
            await update.message.reply_text(results.content)
    else:
        # Handle non-text messages or notify the user accordingly
        print("Received a non-text message.")
        await update.message.reply_text("I can only process text messages.")


async def handle_docs(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    document = update.message.document
    if document.file_name.endswith(('.xlsx', '.xls', '.doc', '.docx')):
        # Get the File object, which contains the file_id
        file = await context.bot.get_file(document.file_id)

        # Construct a full path including the filename
        download_path = os.path.join('/Users/alexanderbauer/Desktop/Cataphora/cityTask/downloads', document.file_name)

        # Use the 'download' method of the 'File' object to save the file to the specified path
        await file.download_to_drive(custom_path=download_path)

        print(f"Document received successfully: {document.file_name}")
        await update.message.reply_text('Document received successfully.')
        if document.file_name.endswith(('.xlsx', '.xls')):
            generate_memo(download_path,
                          f'files/memo.docx',
                          f'files/excel.xlsx')
            await update.message.reply_text('I have successfully generated the memo:')
            print(f"I have successfully generated the memo:")
            await context.bot.send_document(chat_id=update.effective_chat.id,
                                            document=open(f'files/memo.docx', 'rb'))
            await context.bot.send_document(chat_id=update.effective_chat.id,
                                            document=open(f'files/excel.xlsx', 'rb'))
            await context.bot.send_document(chat_id=update.effective_chat.id,
                                            document=open(f'files/Math_Calculations_Work.docx', 'rb'))
    else:
        await update.message.reply_text('Please send an Excel document with the same format as the sample.')


def generate_memo(excel_path, memo_file, workbook_file):
    """
    Modification of the original run.py file to generate a memo. Designed to be used by a telegram bot.
    """
    # Load the Excel file
    filepath = excel_path
    wb = load_workbook(filepath)

    # Load the sheets
    data = wb['Data']
    summary = wb['Summary']
    violation_types = wb['Violation Types']
    print('Sheets Loaded')
    print(violation_types)

    # Search for citations issued by the City Police
    search_text = '01 - CITY POLICE'
    column_to_search = 'B'
    # Convert the column letter to an index (e.g., 'A' -> 1)
    column_index = column_index_from_string(column_to_search)
    index = 0
    # A list to hold your calculations
    math_doc = methods.MathDoc('Math Calculations Work')

    number_of_months = analysis.get_months(summary['A1'].value)
    number_of_months = int(number_of_months.content)

    # Check if 'NOPD Citations' sheet exists, if not create it
    if "NOPD Citations" not in wb.sheetnames:
        nopd_citations = wb.create_sheet("NOPD Citations")
        # Assuming the first row contains headers, copy them first
        headers = [cell.value for cell in data[1]]
        nopd_citations.append(headers)

        # Iterate through all rows in the data sheet
        for row in data.iter_rows(min_row=2):  # Start from row 2 to skip headers
            if row[column_index - 1].value == search_text:  #
                # Check if the cell in your target column matches search_text
                nopd_citations.append([cell.value for cell in row])  # Append the entire row to NOPD Citations sheet
    else:
        nopd_citations = wb["NOPD Citations"]  # Get the existing 'NOPD Citations' sheet

    nopd_data_df = analysis.chart_citations(nopd_citations, 'NOPD')

    # Use ExcelWriter to write to a specific sheet without overwriting other sheets
    with pd.ExcelWriter('files/excel.xlsx', engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
        nopd_data_df.to_excel(writer, sheet_name='NOPD Citations', index=False)

        # Create a dataframe that contains the violation types and their reference codes
        # Set display options to show all columns and rows
        pd.set_option('display.max_columns', None)
        pd.set_option('display.max_rows', None)
        violation_types_df = pd.DataFrame(columns=['violation',
                                                   'local_ref',
                                                   'state_ref',
                                                   'local_fine',
                                                   'state_fine',
                                                   'local_citations',
                                                   'local_citation_revenue',
                                                   'state_citations',
                                                   'state_citation_revenue',
                                                   'nopd_local_citations',
                                                   'nopd_local_citation_revenue',
                                                   'nopd_state_citations',
                                                   'nopd_state_citation_revenue',
                                                   'nopd_total_citations',
                                                   'nopd_total_citation_revenue',
                                                   'nopd_revenue_local_as_state'
                                                   'nopd_total_revenue_local_as_state',
                                                   'nopd_lost_revenue'])

    # Iterate through each row in the worksheet and add the data to the dataframe
    for row in violation_types.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            break
        violation = row[1]  # This is the 'Violation' column
        violation_ref = str(row[0])  # This is the 'Local/State Code Reference' column

        # Determine if the violation is local or state
        if violation_ref.startswith('154'):
            violation_type = 'local_ref'
        else:
            violation_type = 'state_ref'

        # Check if the violation already exists in the dataframe
        if violation in violation_types_df['violation'].values:
            # Update the existing row with the new reference code
            existing_row = violation_types_df[violation_types_df['violation'] == violation].index[0]
            violation_types_df.at[existing_row, violation_type] = violation_ref
        else:
            # Add a new row to the dataframe with the new violation and reference code
            new_row = {'violation': violation, violation_type: violation_ref}
            violation_types_df = violation_types_df._append(new_row, ignore_index=True)

    print(violation_types_df[['violation', 'local_ref', 'state_ref']])

    for index, row in violation_types_df.iterrows():
        # Access Methods that counts the total number of local citations,
        # calculates the revenue generated, writes to the Summary sheet,
        # and updates the dataframe. See methods file for more details
        local_fine, count, rev = methods.total_citations_local(data, summary, row['local_ref'], index)
        violation_types_df.at[index, 'local_fine'] = local_fine
        violation_types_df.at[index, 'local_citations'] = count
        violation_types_df.at[index, 'local_citation_revenue'] = rev

        # Access Methods that counts the total number of state citations,
        # calculates the revenue generated, writes to the Summary sheet,
        # and updates the dataframe. See methods file for more details
        state_fine, count, rev = methods.total_citations_state(data, summary, row['state_ref'], index)
        violation_types_df.at[index, 'state_fine'] = state_fine
        violation_types_df.at[index, 'state_citations'] = count
        violation_types_df.at[index, 'state_citation_revenue'] = rev

        # Access Methods that counts the number of local citations issued by nopd,
        # calculates the revenue generated, writes to the Summary sheet,
        # and updates the dataframe. See methods file for more details
        local_fine, count, rev = methods.nopd_citations_local(nopd_citations, summary, row['local_ref'], index)
        violation_types_df.at[index, 'nopd_local_citations'] = count
        violation_types_df.at[index, 'nopd_local_citation_revenue'] = rev

        # Access Methods that counts the number of state citations issued by nopd,
        # calculates the revenue generated, writes to the Summary sheet,
        # and updates the dataframe. See methods file for more details
        state_fine, count, rev = methods.nopd_citations_state(nopd_citations, summary, row['state_ref'], index)
        violation_types_df.at[index, 'nopd_state_citations'] = count
        violation_types_df.at[index, 'nopd_state_citation_revenue'] = rev

        # Access Methods that calculates the total number of citations issued by NOPD,

        fine_difference, total_citations, total_revenue, citation_difference, \
            total_difference_rev, total_nopd_citations, total_nopd_revenue, nopd_revenue_local_as_state, \
            total_nopd_revenue_local_as_state, nopd_lost_revenue = methods.perform_calculation(math_doc, violation_types_df.at[
                                                                                                            index,
                                                                                                            'violation'
                                                                                                        ],
                                                                           violation_types_df.at[index, 'local_fine'],
                                                                           violation_types_df.at[index, 'state_fine'],
                                                                           violation_types_df.at[
                                                                               index, 'local_citations'],
                                                                           violation_types_df.at[
                                                                               index, 'state_citations'],
                                                                           violation_types_df.at[
                                                                               index, 'local_citation_revenue'],
                                                                           violation_types_df.at[
                                                                               index, 'state_citation_revenue'],
                                                                           violation_types_df.at[
                                                                               index, 'nopd_local_citations'],
                                                                           violation_types_df.at[
                                                                               index, 'nopd_state_citations'],
                                                                           violation_types_df.at[
                                                                               index, 'nopd_local_citation_revenue'],
                                                                           violation_types_df.at[
                                                                               index, 'nopd_state_citation_revenue'])

        # Update the dataframe with the total citations and revenue lost
        violation_types_df.at[index, 'nopd_total_citations'] = total_nopd_citations
        violation_types_df.at[index, 'nopd_total_citation_revenue'] = total_nopd_revenue
        violation_types_df.at[index, 'nopd_revenue_local_as_state'] = nopd_revenue_local_as_state
        violation_types_df.at[index, 'nopd_total_revenue_local_as_state'] = total_nopd_revenue_local_as_state
        violation_types_df.at[index, 'nopd_lost_revenue'] = nopd_lost_revenue
        # Write the lost revenue to the Summary sheet
        methods.write(summary, 'H', index, nopd_lost_revenue)

    print(violation_types_df)

    # Total Columns in the dataframe
    sum_local_citations = violation_types_df['local_citations'].sum()
    sum_local_citation_revenue = violation_types_df['local_citation_revenue'].sum()
    sum_state_citations = violation_types_df['state_citations'].sum()
    sum_state_citation_revenue = violation_types_df['state_citation_revenue'].sum()
    sum_nopd_local_citations = violation_types_df['nopd_local_citations'].sum()
    sum_nopd_local_citation_revenue = violation_types_df['nopd_local_citation_revenue'].sum()
    sum_nopd_state_citations = violation_types_df['nopd_state_citations'].sum()
    sum_nopd_state_citation_revenue = violation_types_df['nopd_state_citation_revenue'].sum()
    sum_nopd_total_citations = violation_types_df['nopd_total_citations'].sum()
    sum_nopd_total_citation_revenue = violation_types_df['nopd_total_citation_revenue'].sum()
    sum_nopd_revenue_local_as_state = violation_types_df['nopd_revenue_local_as_state'].sum()
    sum_total_nopd_revenue_local_as_state = violation_types_df['nopd_total_revenue_local_as_state'].sum()
    sum_nopd_lost_revenue = violation_types_df['nopd_lost_revenue'].sum()

    # Calculate the average monthly lost revenue
    average_monthly_revenue_lost = sum_nopd_lost_revenue / number_of_months
    print(f"\nThe average monthly lost revenue is: {average_monthly_revenue_lost}")
    methods.write(summary, 'B', 14, average_monthly_revenue_lost)

    # Calculate the annualized lost revenue
    annualized_lost_revenue = sum_nopd_lost_revenue / (number_of_months / 12)
    print(f"The annualized lost revenue is: {annualized_lost_revenue}")
    methods.write(summary, 'B', 15, annualized_lost_revenue)

    # Calculate the estimated annual revenue impact
    collection_rate = methods.read_cell_value(summary, 'B19')
    print(f"The collection rate is: {collection_rate}")
    estimated_annual_revenue_impact = annualized_lost_revenue * collection_rate
    print(f"The estimated annual revenue impact is: {estimated_annual_revenue_impact}\n")
    methods.write(summary, 'B', 17, estimated_annual_revenue_impact)

    # Create a DataFrame with the calculation results
    data = {
        'Description': ['Average Monthly Lost Revenue', 'Annualized Lost Revenue', 'Collection Rate',
                        'Estimated Annual Revenue Impact'],
        'Value': [average_monthly_revenue_lost, annualized_lost_revenue, collection_rate,
                  estimated_annual_revenue_impact]
    }
    results_df = pd.DataFrame(data)

    # Update the dataframe with the totals
    totals_df = pd.DataFrame(columns=[
        'sum_local_citations', 'sum_local_citation_revenue',
        'sum_state_citations', 'sum_state_citation_revenue',
        'sum_nopd_local_citations', 'sum_nopd_local_citation_revenue',
        'sum_nopd_state_citations', 'sum_nopd_state_citation_revenue',
        'sum_nopd_total_citations', 'sum_nopd_total_citation_revenue',
        'sum_nopd_revenue_local_as_state',
        'sum_nopd_total_revenue_local_as_state', 'sum_nopd_lost_revenue'])
    totals_df.at[0, 'sum_local_citations'] = sum_local_citations
    totals_df.at[0, 'sum_local_citation_revenue'] = sum_local_citation_revenue
    totals_df.at[0, 'sum_state_citations'] = sum_state_citations
    totals_df.at[0, 'sum_state_citation_revenue'] = sum_state_citation_revenue
    totals_df.at[0, 'sum_nopd_local_citations'] = sum_nopd_local_citations
    totals_df.at[0, 'sum_nopd_local_citation_revenue'] = sum_nopd_local_citation_revenue
    totals_df.at[0, 'sum_nopd_state_citations'] = sum_nopd_state_citations
    totals_df.at[0, 'sum_nopd_state_citation_revenue'] = sum_nopd_state_citation_revenue
    totals_df.at[0, 'sum_nopd_total_citations'] = sum_nopd_total_citations
    totals_df.at[0, 'sum_nopd_total_citation_revenue'] = sum_nopd_total_citation_revenue
    totals_df.at[0, 'sum_nopd_revenue_local_as_state'] = sum_nopd_revenue_local_as_state
    totals_df.at[0, 'sum_total_nopd_revenue_local_as_state'] = sum_total_nopd_revenue_local_as_state
    totals_df.at[0, 'sum_nopd_lost_revenue'] = sum_nopd_lost_revenue

    # Transpose the DataFrame
    totals_df_transposed = totals_df.T

    # you need to reset the index so that it becomes a regular column in the DataFrame
    totals_df_transposed.reset_index(inplace=True)

    # rename the columns to reflect the change
    totals_df_transposed.columns = ['Type of Total', 'Value']

    methods.MathDoc.add_dataframe(math_doc, totals_df_transposed)
    methods.perform_final_calculation(math_doc, sum_nopd_lost_revenue, number_of_months, collection_rate)

    # DataFrame for citations
    citations_df_by_violation = violation_types_df[['violation',
                                                    'nopd_local_citations', 'nopd_state_citations',
                                                    'nopd_total_citations']]

    # Rename the columns as needed
    citations_df_by_violation = citations_df_by_violation.rename(columns={
        'violation': 'Violation',
        'nopd_local_citations': 'NOPD Local Citations',
        'nopd_state_citations': 'NOPD State Citations',
        'nopd_total_citations': 'NOPD Total Citations'})

    # Create a DataFrame for total citations
    citations_df = totals_df[['sum_nopd_local_citations', 'sum_nopd_state_citations',
                              'sum_nopd_total_citations']]

    # Rename the columns as needed
    citations_df = citations_df.rename(columns={
        'sum_nopd_local_citations': 'NOPD Local Citations',
        'sum_nopd_state_citations': 'NOPD State Citations',
        'sum_nopd_total_citations': 'NOPD Total Citations'
        # Add more columns as needed
    })

    # Extract the totals from `citations_df`
    citation_totals_data = citations_df.iloc[0]  # This gets the first (and only) row of totals

    # Create a new DataFrame for the totals row, with 'Totals' as the violation
    # Make sure the columns match the `citations_df_by_violation` DataFrame
    citation_totals_row = pd.DataFrame([['Total NOPD Citations'] + citation_totals_data.values.tolist()],
                                       columns=citations_df_by_violation.columns)

    # Append the totals row to the `citations_df_by_violation` DataFrame
    citations_df_by_violation = citations_df_by_violation._append(citation_totals_row, ignore_index=True)

    # DataFrame for revenue
    revenue_df_by_violation = violation_types_df[['violation',
                                                  'nopd_local_citation_revenue', 'nopd_state_citation_revenue',
                                                  'nopd_total_citation_revenue', 'nopd_revenue_local_as_state',
                                                  'nopd_total_revenue_local_as_state', 'nopd_lost_revenue']]

    # Create a DataFrame for total revenue
    revenue_df = totals_df[['sum_nopd_local_citation_revenue', 'sum_nopd_state_citation_revenue',
                            'sum_nopd_total_citation_revenue', 'sum_nopd_revenue_local_as_state',
                            'sum_total_nopd_revenue_local_as_state', 'sum_nopd_lost_revenue']]

    # Rename the columns as needed
    revenue_df = revenue_df.rename(columns={
        'sum_nopd_local_citation_revenue': 'NOPD Local Revenue',
        'sum_nopd_state_citation_revenue': 'NOPD State Revenue',
        'sum_nopd_total_citation_revenue': 'NOPD Total Revenue',
        'sum_nopd_revenue_local_as_state': 'NOPD Revenue Local as State',
        'sum_total_nopd_revenue_local_as_state': 'Total NOPD Revenue Local as State',
        'sum_nopd_lost_revenue': 'NOPD Lost Revenue'
        # Add more columns as needed
    })

    # Extract the totals from `citations_df`
    revenue_totals_data = revenue_df.iloc[0]  # This gets the first (and only) row of totals

    # Create a new DataFrame for the totals row, with 'Totals' as the violation
    # Make sure the columns match the `revenue_df_by_violation` DataFrame
    revenue_totals_row = pd.DataFrame([['Total NOPD Revenue'] + revenue_totals_data.values.tolist()],
                                      columns=revenue_df_by_violation.columns)

    # Append the totals row to the `citations_df_by_violation` DataFrame
    revenue_df_by_violation = revenue_df_by_violation._append(revenue_totals_row, ignore_index=True)

    # Rename the columns as needed
    revenue_df_by_violation = revenue_df_by_violation.rename(columns={
        'violation': 'Violation',
        'nopd_local_citation_revenue': 'NOPD Local Revenue',
        'nopd_state_citation_revenue': 'NOPD State Revenue',
        'nopd_total_citation_revenue': 'NOPD Total Revenue',
        'nopd_revenue_local_as_state': 'NOPD Revenue Local as State',
        'nopd_total_revenue_local_as_state': 'Total NOPD Revenue Local as State',
        'nopd_lost_revenue': 'NOPD Lost Revenue'})

    # Write the totals to the Summary sheet
    methods.write(summary, 'D', 11, sum_local_citations)
    methods.write(summary, 'E', 11, sum_state_citations)
    methods.write(summary, 'F', 11, sum_nopd_local_citations)
    methods.write(summary, 'G', 11, sum_nopd_state_citations)
    methods.write(summary, 'H', 11, sum_nopd_lost_revenue)

    # Load the template Word document
    doc = Document('files/For Review Memo Template.docx')

    try:
        memo_title = analysis.get_memo_completion('Write a 3 to 5 word title for the memo',
                                              violation_types_df, totals_df, results_df, analysis.gpt4, '')
    except:
        memo_title = analysis.get_memo_completion('Write a 3 to 5 word title for the memo',
                                              violation_types_df, totals_df, results_df, analysis.gpt3, '')

    generated_text = 'Title:' + memo_title.content

    # Access and edit the content of the document
    for paragraph in doc.paragraphs:
        if 'To:' in paragraph.text:
            new_run = paragraph.add_run('Gilbert Montaño, ')
            new_run.bold = False
            new_run.font.size = Pt(11)
            new_run.font.name = 'Arial'
            new_run = paragraph.add_run('Chief Administrative Officer')
            new_run.italic = True
            new_run.font.size = Pt(11)
            new_run.font.name = 'Arial'
        elif 'From:' in paragraph.text:
            new_run = paragraph.add_run('Alex Bauer')
            new_run.bold = False
            new_run.font.size = Pt(11)
            new_run.font.name = 'Arial'
        elif 'CC:' in paragraph.text:
            new_run = paragraph.add_run('Cameron MacPhee, Rebecca Houtman')
            new_run.bold = False
            new_run.font.size = Pt(11)
            new_run.font.name = 'Arial'
        elif 'Date:' in paragraph.text:
            new_run = paragraph.add_run('February 2, 2024')
            new_run.bold = False
            new_run.font.size = Pt(11)
            new_run.font.name = 'Arial'
        elif 'FOR REVIEW:' in paragraph.text:
            new_run = paragraph.add_run(memo_title.content)
            new_run.bold = False
            new_run.font.size = Pt(11)
            new_run.font.name = 'Arial'

    # Add a bold header
    background_heading = doc.add_paragraph('BACKGROUND')
    # Set spacing before "BACKGROUND" to zero
    background_heading.paragraph_format.space_before = Pt(0)
    background_heading.runs[0].font.bold = True
    background_heading.runs[0].font.size = Pt(11)
    background_heading.runs[0].font.name = 'Arial'

    try:
        memo_background = analysis.get_memo_completion('Write a background section for the memo, '
                                                       'no need to label this secton as "BACKGROUND" as '
                                                       'it is already clear from the heading. ',
                                                       violation_types_df, totals_df, results_df,
                                                       analysis.gpt4, generated_text)
    except:
        logger.error('GPT-4 failed to generate the memo background section. Using GPT-3 instead.')
        memo_background = analysis.get_memo_completion('Write a background section for the memo, '
                                                       'no need to label this secton as "BACKGROUND" as '
                                                       'it is already clear from the heading. ',
                                                       violation_types_df, totals_df, results_df,
                                                       analysis.gpt3, generated_text)

    generated_text = generated_text + 'Background: ' + memo_background.content

    # Add a paragraph underneath the header
    background_text = doc.add_paragraph(memo_background.content)
    background_text.runs[0].font.name = 'Arial'

    # Add another heading
    task_heading = doc.add_paragraph('\nANALYSIS')
    task_heading.runs[0].font.bold = True
    task_heading.runs[0].font.size = Pt(11)
    task_heading.runs[0].font.name = 'Arial'

    try:
        memo_analysis = analysis.get_memo_completion('Write an introduction to the analysis section for the memo. '
                                                     'This should be one to two paragraphs long. '
                                                     'no need to label this section as "ANALYSIS" as '
                                                     'it is already clear from the heading'
                                                     'Discuss how there is multiple ways to measure lost revenue.'
                                                     'However for this analysis we will be focusing on how much more revenue '
                                                     'NOPD could have brought in for the City of New Orleans if they decided to '
                                                     'cite drivers for state violations instead of local ordinaces. '
                                                     'Lost revenue will be defined as the difference between the actual revenue '
                                                     'NOPD brought in and the potential revenue they could have brought in '
                                                     'if all citations were state violations. '
                                                     'Lay out the key assumptions of analyzing the data in this way.',
                                                     violation_types_df, totals_df, results_df,
                                                     analysis.gpt4, generated_text)
    except:
        logger.error('GPT-4 failed to generate the memo analysis section. Using GPT-3 instead.')
        memo_analysis = analysis.get_memo_completion('Write an introduction to the analysis section for the memo. '
                                                     'This should be one to two paragraphs long. '
                                                     'no need to label this section as "ANALYSIS" as '
                                                     'it is already clear from the heading'
                                                     'Discuss how there is multiple ways to measure lost revenue.'
                                                     'However for this analysis we will be focusing on how much more revenue '
                                                     'NOPD could have brought in for the City of New Orleans if they decided to '
                                                     'cite drivers for state violations instead of local ordinaces. '
                                                     'Lost revenue will be defined as the difference between the actual revenue '
                                                     'NOPD brought in and the potential revenue they could have brought in '
                                                     'if all citations were state violations. '
                                                     'Lay out the key assumptions of analyzing the data in this way.',
                                                     violation_types_df, totals_df, results_df,
                                                     analysis.gpt3, generated_text)

    generated_text = generated_text + 'Analysis: ' + memo_analysis.content

    # Add another paragraph underneath the second heading
    task_text = doc.add_paragraph(memo_analysis.content)
    task_text.runs[0].font.name = 'Arial'

    # Add another heading
    findings_heading = doc.add_paragraph('\nKEY FINDINGS')
    findings_heading.runs[0].font.bold = True
    findings_heading.runs[0].font.size = Pt(11)
    findings_heading.runs[0].font.name = 'Arial'

    # Add Tables to the Word Document. See Methods.py for the write_table method
    methods.write_table(doc, citations_df_by_violation, False, 'Citations by Violation Type')
    methods.write_table(doc, citations_df, False, 'Total Citations')

    doc.add_paragraph()
    doc.add_picture('files/NOPD_citations_plot.png', width=Inches(6))

    try:
        memo_citation_table = analysis.get_memo_completion('Briefly summarize only the citation and total citation data, '
                                                           'without discussing any revenue implications in one paragraph.'
                                                           'discuss how the table demonstrates majority of citations are for local violations'
                                                           'except for brake tag violations, which are mostly state violations.'
                                                           'discuss how besides from brake tag violations, '
                                                           'most of the local violations came from careless operation of a vehicle,'
                                                           'drivers to use reasonable vigilance, and following too closely,'
                                                           'which all are somewhat up for interpretation by NOPD officers.',
                                                           violation_types_df, totals_df, results_df,
                                                           analysis.gpt4, generated_text)
    except:
        logger.error('GPT-4 failed to generate the memo citation table section. Using GPT-3 instead.')
        memo_citation_table = analysis.get_memo_completion('Briefly summarize only the citation and total citation data, '
                                                           'without discussing any revenue implications in one paragraph.'
                                                           'discuss how the table demonstrates majority of citations are for local violations'
                                                           'except for brake tag violations, which are mostly state violations.'
                                                           'discuss how besides from brake tag violations, '
                                                           'most of the local violations came from careless operation of a vehicle,'
                                                           'drivers to use reasonable vigilance, and following too closely,'
                                                           'which all are somewhat up for interpretation by NOPD officers.',
                                                           violation_types_df, totals_df, results_df,
                                                           analysis.gpt3, generated_text)

    generated_text = generated_text + 'Citation Table: ' + memo_citation_table.content

    background_text = doc.add_paragraph(memo_citation_table.content)
    background_text.runs[0].font.name = 'Arial'

    # Add a blank paragraph, which will appear as a blank line in the document
    doc.add_paragraph()

    # Add Tables to the Word Document. See Methods.py for the write_table method
    methods.write_table(doc, revenue_df_by_violation, True, 'Revenue by Violation Type')
    methods.write_table(doc, revenue_df, True, 'Total Revenue')

    doc.add_paragraph()
    doc.add_picture('files/NOPD_revenue_plot.png', width=Inches(6))

    try:
        memo_revenue_table = analysis.get_memo_completion('Briefly summarize only the revenue and total revenue data'
                                                          'in one paragraph. discuss how the table demonstrates'
                                                          'that the majority of revenue comes from unlawful use of licence violations,'
                                                          'which are one of the most expensive violations. '
                                                          'discuss how the table also demonstrates that the area with the most potential'
                                                          'revenue is license plate violations, which are almost all local violations,'
                                                          'however the state fine is almost double the local fine ($145 more than $157.50). ',
                                                          violation_types_df, totals_df, results_df,
                                                          analysis.gpt4, generated_text)
    except:
        logger.error('GPT-4 failed to generate the memo revenue table section. Using GPT-3 instead.')
        memo_revenue_table = analysis.get_memo_completion('Briefly summarize only the revenue and total revenue data'
                                                          'in one paragraph. discuss how the table demonstrates'
                                                          'that the majority of revenue comes from unlawful use of licence violations,'
                                                          'which are one of the most expensive violations. '
                                                          'discuss how the table also demonstrates that the area with the most potential'
                                                          'revenue is license plate violations, which are almost all local violations,'
                                                          'however the state fine is almost double the local fine ($145 more than $157.50). ',
                                                          violation_types_df, totals_df, results_df,
                                                          analysis.gpt3, generated_text)

    generated_text = generated_text + 'Revenue Table: ' + memo_revenue_table.content

    background_text = doc.add_paragraph(memo_revenue_table.content)
    background_text.runs[0].font.name = 'Arial'

    # Add a blank paragraph, which will appear as a blank line in the document
    doc.add_paragraph()

    # Call your function to add the table to the document
    methods.write_table(doc, results_df, money=True, title="Financial Impact Analysis")

    try:
        memo_impact_table = analysis.get_memo_completion('Summarize the financial impact analysis in results_df. '
                                                         'include a discussion of the average monthly lost revenue potential, '
                                                         'the annualized lost revenue potential, '
                                                         'the total lost revenue potential, '
                                                         'and the estimated annual revenue impact. '
                                                         'Discuss the implications of the results, especially focusing on the '
                                                         'areas where revenue was lost and how the NOPD could improve. '
                                                         'make sure to include the dollar amounts in the summary, '
                                                         'and respond in 3 paragraphs without using bullet points or lists.',
                                                         violation_types_df, totals_df, results_df,
                                                         analysis.gpt4, generated_text)
    except:
        logger.error('GPT-4 failed to generate the memo impact table section. Using GPT-3 instead.')
        memo_impact_table = analysis.get_memo_completion('Summarize the financial impact analysis in results_df. '
                                                         'include a discussion of the average monthly lost revenue potential, '
                                                         'the annualized lost revenue potential, '
                                                         'the total lost revenue potential, '
                                                         'and the estimated annual revenue impact. '
                                                         'Discuss the implications of the results, especially focusing on the '
                                                         'areas where revenue was lost and how the NOPD could improve. '
                                                         'make sure to include the dollar amounts in the summary, '
                                                         'and respond in 3 paragraphs without using bullet points or lists.',
                                                         violation_types_df, totals_df, results_df,
                                                         analysis.gpt3, generated_text)

    generated_text = generated_text + 'Impact Table: ' + memo_impact_table.content

    background_text = doc.add_paragraph(memo_impact_table.content)
    background_text.runs[0].font.name = 'Arial'

    # Add a blank paragraph, which will appear as a blank line in the document
    doc.add_paragraph()

    # Add another heading
    task_heading = doc.add_paragraph('RECOMMENDATION')
    task_heading.runs[0].font.bold = True
    task_heading.runs[0].font.size = Pt(11)
    task_heading.runs[0].font.name = 'Arial'

    try:
        memo_recommendation = analysis.get_memo_completion('Write a 4-5 paragraph recommendation section that discusses '
                                                           'the implications of '
                                                           'the results of the analysis, '
                                                           'and what the next steps should be. '
                                                           'This should be a brief summary '
                                                           'of the implications of the analysis.'
                                                           'no need to label this secton as "RECOMMENDATION" as '
                                                           'it is already clear from the heading. '
                                                           'when writing the recommendation, '
                                                           'make sure to include the dollar amounts in the summary, '
                                                           'and bring up the estimated annual revenue impact. '
                                                           'discuss how upon discussing the analysis with a '
                                                           'NOPD traffic citation investigator, they revealed that '
                                                           'the NOPD uses state statutes to issue citations, when they '
                                                           'feel like the violation is more serious, it is their way of '
                                                           'determining who gets punished more or '
                                                           'less for their infractions',
                                                           violation_types_df, totals_df, results_df,
                                                           analysis.gpt4, generated_text)
    except:
        logger.error('GPT-4 failed to generate the memo recommendation section. Using GPT-3 instead.')
        memo_recommendation = analysis.get_memo_completion('Write a 4-5 paragraph recommendation section that discusses '
                                                           'the implications of '
                                                           'the results of the analysis, '
                                                           'and what the next steps should be. '
                                                           'This should be a brief summary '
                                                           'of the implications of the analysis.'
                                                           'no need to label this secton as "RECOMMENDATION" as '
                                                           'it is already clear from the heading. '
                                                           'when writing the recommendation, '
                                                           'make sure to include the dollar amounts in the summary, '
                                                           'and bring up the estimated annual revenue impact. '
                                                           'discuss how upon discussing the analysis with a '
                                                           'NOPD traffic citation investigator, they revealed that '
                                                           'the NOPD uses state statutes to issue citations, when they '
                                                           'feel like the violation is more serious, it is their way of '
                                                           'determining who gets punished more or '
                                                           'less for their infractions',
                                                           violation_types_df, totals_df, results_df,
                                                           analysis.gpt3, generated_text)

    generated_text = generated_text + 'Recommendation: ' + memo_recommendation.content

    print(generated_text)

    # Add another paragraph underneath the second heading
    task_text = doc.add_paragraph(memo_recommendation.content)
    task_text.runs[0].font.name = 'Arial'

    # Save the document after all calculations are added
    math_doc.save_doc('files/Math_Calculations_Work.docx')

    # Save the changes to the Excel file
    wb.save(workbook_file)
    # Save the changes to the Word document
    doc.save(memo_file)

    # Close the workbook
    wb.close()

    print('Memo has been created and saved to the following location: ', filepath)

    return


def main() -> None:
    """Start the bot."""
    # Create the Application and pass it your bot's token.
    application = Application.builder().token(telegram_api_key).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("get_memo", get_memo))
    application.add_handler(MessageHandler(filters.TEXT, handle_message))

    document_handler = MessageHandler(filters.Document.ALL, handle_docs)
    application.add_handler(document_handler)

    # Run the bot until the user presses Ctrl-C
    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
