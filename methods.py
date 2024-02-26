from openpyxl.utils import column_index_from_string

from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor


class MathDoc:
    """
    A class to represent a Word document with mathematical calculations
    """
    def __init__(self, doc_title):
        # Initialize a new Word document
        self.doc = Document()
        self.doc.add_heading(doc_title, level=1)

    def add_math_calculation(self, calculation):
        # Add a paragraph with the calculation
        self.doc.add_paragraph(calculation)

    def add_table(self, df, money, title):
        write_table(self.doc, df, money, title)

    def add_dataframe(self, df):
        # Add a dataframe to the document
        table = self.doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = 'Table Grid'
        for i in range(df.shape[1]):
            # Convert header to string to ensure compatibility
            table.cell(0, i).text = str(df.columns[i])
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                # Convert each cell's value to string to ensure compatibility
                table.cell(i + 1, j).text = str(df.iloc[i, j])

    def save_doc(self, file_name):
        # Save the document to the given file name
        self.doc.save(file_name)


def perform_calculation(math_doc, violation, local_fine, state_fine, local, state, local_rev, state_rev,
                        nopd_local, nopd_state, nopd_local_rev, nopd_state_rev):
    """
    Perform a series of calculations and add them to the math_doc
    :param math_doc:
    :param violation:
    :param local_fine:
    :param state_fine:
    :param local:
    :param state:
    :param local_rev:
    :param state_rev:
    :param nopd_local:
    :param nopd_state:
    :param nopd_local_rev:
    :param nopd_state_rev:
    :return:
    """
    MathDoc.add_math_calculation(math_doc, f"Violation: {violation}")

    fine_difference = state_fine - local_fine
    MathDoc.add_math_calculation(math_doc, f"Fine Difference: \n"
                                           f"State Fine - Local Fine = Fine Difference\n "
                                           f"{state_fine} - {local_fine} = {fine_difference}\n"
                                           f"The state fine is {fine_difference} more than the local fine.\n\n")

    total_citations = local + state
    MathDoc.add_math_calculation(math_doc, f"Total Citations (All Dept): \n"
                                           f"Sum(Local Citations) + Sum(State Citations) = Total Citations\n "
                                           f"{local} + {state} = {total_citations}\n\n")

    total_revenue = local_rev + state_rev
    MathDoc.add_math_calculation(math_doc, f"Total Revenue (All Dept): \n"
                                           f"Sum(Local Revenue) + Sum(State Revenue) = Total Revenue\n "
                                           f"{local_rev} + {state_rev} = {total_revenue}\n\n")

    citation_difference = local - state
    MathDoc.add_math_calculation(math_doc, f"Local Citations that could have been State Citations: \n"
                                           f"Sum(Local Citations) - Sum(State Citations) = Local Citations that could have been State Citations\n "
                                           f"{local} - {state} = {citation_difference}\n")

    total_difference_rev = citation_difference * fine_difference
    MathDoc.add_math_calculation(math_doc, f"Local Revenue that could have been State Revenue: \n"
                                           f"Local Citations that could have been State Citations * Fine Difference = Local Revenue that could have been State Revenue\n "
                                           f"{citation_difference} * {fine_difference} = {total_difference_rev}\n")

    total_nopd_citations = nopd_local + nopd_state
    MathDoc.add_math_calculation(math_doc, f"Total NOPD Citations: \n"
                                           f"Sum NOPD Local Citations + Sum NOPD State Citations = Total NOPD Citations\n "
                                           f"{nopd_local} + {nopd_state} = {total_nopd_citations}\n\n")

    total_nopd_revenue = nopd_local_rev + nopd_state_rev
    MathDoc.add_math_calculation(math_doc, f"Total NOPD Revenue: \n"  
                                           f"Sum NOPD Local Revenue + Sum NOPD State Revenue = Total NOPD Revenue\n "
                                           f"{nopd_local_rev} + {nopd_state_rev} = {total_nopd_revenue}\n")

    nopd_revenue_local_as_state = nopd_local * state_fine
    MathDoc.add_math_calculation(math_doc, f"NOPD Local Citations as State Citations: \n"
                                             f"Sum NOPD Local Citations * State Fine = NOPD Local Citations as State Citations\n "
                                           f" State ({nopd_local} * {state_fine}) = {nopd_revenue_local_as_state}\n")

    total_nopd_revenue_local_as_state = nopd_state_rev + (nopd_local * state_fine)
    MathDoc.add_math_calculation(math_doc, f"NOPD Local Citations as State Citations: \n"
                                           f"Sum NOPD Local Citations * State Fine = NOPD Local Citations as State Citations\n "
                                           f" State ({nopd_local} * {state_fine}) = {total_nopd_revenue_local_as_state}\n")

    nopd_lost_revenue = nopd_revenue_local_as_state - nopd_local_rev
    MathDoc.add_math_calculation(math_doc, f"NOPD Lost Revenue: \n"
                                           f"NOPD Local Citations as State Citations - NOPD Local Revenue = NOPD Lost Revenue\n "
                                           f"{nopd_revenue_local_as_state} - {nopd_local_rev} = {nopd_lost_revenue}\n")

    return fine_difference, total_citations, total_revenue, citation_difference, total_difference_rev, \
        total_nopd_citations, total_nopd_revenue, nopd_revenue_local_as_state, total_nopd_revenue_local_as_state, \
        nopd_lost_revenue


def perform_final_calculation(math_doc, sum_nopd_lost_revenue, number_of_months, collection_rate):
    """
    Perform a series of calculations and add them to the math_doc
    :param math_doc:
    :param sum_nopd_lost_revenue:
    :param number_of_months:
    :param collection_rate:
    :return:
    """
    MathDoc.add_math_calculation(math_doc, f"Revenue Impact Calculations: \n\n")
    average_monthly_revenue_lost = sum_nopd_lost_revenue / number_of_months
    MathDoc.add_math_calculation(math_doc, f"Average Monthly Revenue Lost: \n"
                                           f"Sum NOPD Lost Revenue / Number of Months = Average Monthly Revenue Lost\n "
                                           f"{sum_nopd_lost_revenue} / {number_of_months} = {average_monthly_revenue_lost}\n\n")
    annualized_lost_revenue = sum_nopd_lost_revenue / (number_of_months / 12)
    MathDoc.add_math_calculation(math_doc, f"Annualized Lost Revenue: \n"
                                           f"Sum NOPD Lost Revenue / (Number of Months / 12) = Annualized Lost Revenue\n "
                                           f"{sum_nopd_lost_revenue} / ({number_of_months} / 12) = {annualized_lost_revenue}\n\n")
    estimated_annual_revenue_impact = annualized_lost_revenue * collection_rate
    MathDoc.add_math_calculation(math_doc, f"Estimated Annual Revenue Impact: \n"
                                           f"Annualized Lost Revenue * Collection Rate = Estimated Annual Revenue Impact\n "
                                           f"{annualized_lost_revenue} * {collection_rate} = {estimated_annual_revenue_impact}\n\n")
    MathDoc.add_math_calculation(math_doc, f"-------------------\n\n\n")


def set_cell_borders(cell, **kwargs):
    """
    Set cell's border color, thickness, and other attributes
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Set the cell borders
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = parse_xml(r'<w:tcBorders {}><{}/></w:tcBorders>'.format(nsdecls('w'), tag))
            element[0].set(parse_xml(r'<w:sz w:val="{}"/>'.format(edge_data['sz']))[0].tag, edge_data['sz'])
            element[0].set(parse_xml(r'<w:color w:val="{}"/>'.format(edge_data['color']))[0].tag, edge_data['color'])
            tcPr.append(element)

# Set the background color of a cell
def set_cell_background_color(cell, color_str):
    """
    Set the background color of a cell.
    """
    # Create a new shading element
    shd = OxmlElement('w:shd')
    # Set the attributes for the shading element
    shd.set(qn('w:fill'), color_str)

    # Get the cell properties element and add the shading element to it
    cell._tc.get_or_add_tcPr().append(shd)

def count_rows_with_text(worksheet, column_letter, search_text):
    """
    Count the number of rows in a column that contain a specific text
    :param worksheet:
    :param column_letter:
    :param search_text:
    :return:
    """
    # Convert the column letter to an index (e.g., 'A' -> 1)
    column_index = column_index_from_string(column_letter)

    # Initialize a counter for the rows containing the search text
    count = 0

    # Iterate through all rows in the specified column
    for row in worksheet.iter_rows(min_col=column_index, max_col=column_index):
        cell = row[0]  # Since we're looking at one column, there will only be one cell per row
        if cell.value == search_text:
            count += 1

    return count

def read_cell_value(sheet, cell_address):
    """
    Read the value of a cell
    :param sheet:
    :param cell_address:
    :return:
    """
    # Check if the cell contains a value
    cell = sheet[cell_address]  # Replace with the specific cell address
    if cell.value is not None:
        # Write to the cell
        return cell.value
    else:
        return None

def write_cell_value(sheet, cell_address, value):
    """
    Write a value to a cell
    :param sheet:
    :param cell_address:
    :param value:
    :return:
    """
    # Check if the cell contains a value
    cell = sheet[cell_address]  # Replace with the specific cell address
    if cell.value is None:
        # Write to the cell
        cell.value = value
        print(f"The cell {cell.coordinate} has been updated to {value}.")


def total_citations_local(data, summary, violation_ref, index):
    """
    Count the total number of local citations, calculate revenue, and write to the Summary sheet
    :param data:
    :param summary:
    :param violation_ref:
    :param index:
    :return:
    """
    # Count the total number of local citations, calculate revenue, and write to the Summary sheet
    # index in the summary sheet is shifted due to headers
    index = index + 3

    local_fine = summary.cell(row=index, column=2).value

    # Count the total number of local citations
    total_local_citation_count = count_rows_with_text(data, 'D', violation_ref)
    # Write the total number of local citations to the Summary sheet
    write_cell_value(summary, f'D{index}', total_local_citation_count)

    # Calculate the revenue from local Citations
    total_local_citation_rev = total_local_citation_count * local_fine
    # Return the total number of local citations and the total revenue
    return local_fine, total_local_citation_count, total_local_citation_rev

def total_citations_state(data, summary, violation_ref, index):
    """
    Count the total number of local citations, calculate revenue, and write to the Summary sheet
    :param data:
    :param summary:
    :param violation_ref:
    :param index:
    :return:
    """
    # Count the total number of local citations, calculate revenue, and write to the Summary sheet
    # index in the summary sheet is shifted due to headers
    index = index + 3

    state_fine = summary.cell(row=index, column=3).value

    # Count the total number of local citations
    total_state_citation_count = count_rows_with_text(data, 'D', violation_ref)
    # Write the total number of local citations to the Summary sheet
    write_cell_value(summary, f'F{index}', total_state_citation_count)

    # Calculate the revenue from local Citations
    total_state_citation_rev = total_state_citation_count * state_fine
    # Return the total number of local citations and the total revenue
    return state_fine, total_state_citation_count, total_state_citation_rev


def nopd_citations_local(data, summary, violation_ref, index):
    """
    Count the total number of local citations, calculate revenue, and write to the Summary sheet
    :param data:
    :param summary:
    :param violation_ref:
    :param index:
    :return:
    """
    # Count the total number of local citations, calculate revenue, and write to the Summary sheet
    # index in the summary sheet is shifted due to headers
    index = index + 3

    local_fine = summary.cell(row=index, column=2).value

    # Count the total number of local citations
    total_local_citation_count = count_rows_with_text(data, 'D', violation_ref)
    # Write the total number of local citations to the Summary sheet
    write_cell_value(summary, f'E{index}', total_local_citation_count)

    # Calculate the revenue from local Citations
    total_local_citation_rev = total_local_citation_count * summary[f'B{index}'].value
    # Return the total number of local citations and the total revenue
    return local_fine, total_local_citation_count, total_local_citation_rev


def nopd_citations_state(data, summary, violation_ref, index):
    """
    Count the total number of local citations, calculate revenue, and write to the Summary sheet
    :param data:
    :param summary:
    :param violation_ref:
    :param index:
    :return:
    """
    # Count the total number of local citations, calculate revenue, and write to the Summary sheet
    # index in the summary sheet is shifted due to headers
    index = index + 3

    state_fine = summary.cell(row=index, column=3).value

    # Count the total number of local citations
    total_state_citation_count = count_rows_with_text(data, 'D', violation_ref)
    # Write the total number of local citations to the Summary sheet
    write_cell_value(summary, f'G{index}', total_state_citation_count)

    # Calculate the revenue from local Citations
    total_state_citation_rev = total_state_citation_count * summary[f'C{index}'].value
    # Return the total number of local citations and the total revenue
    return state_fine, total_state_citation_count, total_state_citation_rev

def write_lost_rev(sheet, index, value):
    """
    Write the lost revenue to the Summary sheet
    :param sheet:
    :param index:
    :param value:
    :return:
    """
    # index in the summary sheet is shifted due to headers
    index = index + 3
    # Write the total number of local citations to the Summary sheet
    write_cell_value(sheet, f'H{index}', value)

def write(sheet, column, index, value):
    """
    Write the lost revenue to the Summary sheet
    :param sheet:
    :param column:
    :param index:
    :param value:
    :return:
    """
    # index in the summary sheet is shifted due to headers
    index = index + 3
    # Write the total number of local citations to the Summary sheet
    write_cell_value(sheet, f'{column}{index}', value)
    
def write_table(doc, df, money, title):
    """
    Add a table to the Word document with an additional row for headers
    :param doc:
    :param df:
    :param money:
    :param title:
    :return:
    """
    # Add a table to the Word document with an additional row for headers
    table = doc.add_table(df.shape[0] + 2, df.shape[1])

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set autofit to false so we can set widths as needed
    table.autofit = False

    # Define the style for the table - you can customize this as needed
    table.style = 'Table Grid'

    # Merge the cells in the first row to create a title cell
    title_cell = table.cell(0, 0)
    title_cell.merge(table.cell(0, df.shape[1] - 1))
    title_cell.text = f"{title}"
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment

    # Set the background color of the title cell to dark blue
    shading_elm = parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w')))
    title_cell._tc.get_or_add_tcPr().append(shading_elm)

    # Set the font of the title cell to bold and white
    for paragraph in title_cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True  # Set the font to bold
            run.font.color.rgb = RGBColor(255, 255, 255)  # Set the font color to white

    # Add the column headers from the DataFrame and format them for readability
    for i, column in enumerate(df.columns):
        cell = table.cell(1, i)
        cell.text = column
        paragraph = cell.paragraphs[0]
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)  # Set font size for readability
        run.font.name = 'Arial'  # Choose a clear and readable font type
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment
        cell.width = Inches(1)  # Set the width of the cell

        # Set the background color to a light shade for contrast (e.g., light gray)
        shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Set the header row formatting (dark blue background with white text)
    for i, cell in enumerate(table.rows[0].cells):
        # Set the font color to white and bold
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True

        # Set the background color to dark blue
        shading_elm = parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Add the data from the DataFrame to the table with dollar signs
    for i in range(df.shape[0]):  # Loop over rows
        for j in range(df.shape[1]):  # Loop over columns
            cell = table.cell(i + 2, j)  # '+2' to account for title and header rows
            value = df.iloc[i, j]
            # Check if the value is a number and should be formatted as currency
            if money == True:
                if isinstance(value, (int, float)):
                    cell.text = f"${value:,.2f}"  # Format the number as currency with 2 decimal places
                else:
                    cell.text = str(value)
            else:
                cell.text = str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center alignment

